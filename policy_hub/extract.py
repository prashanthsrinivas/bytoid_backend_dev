"""Text extraction helpers for Policy Hub uploads.

Converts raw uploaded bytes (PDF / DOCX / HTML) into a sanitised HTML
fragment suitable for downstream LLM mapping into V2 structured sections.
"""

import io
import os
import re
from html import escape

from bs4 import BeautifulSoup

from utils.base_logger import get_logger

logger = get_logger(__name__)


_DANGEROUS_TAGS = {"script", "style", "iframe", "form", "object", "embed", "meta", "link"}


def _paragraphs_from_text(text: str) -> list[str]:
    blocks = re.split(r"\n\s*\n", text)
    parts: list[str] = []
    for blk in blocks:
        cleaned = " ".join(line.strip() for line in blk.splitlines() if line.strip())
        if cleaned:
            parts.append(f"<p>{escape(cleaned)}</p>")
    return parts


def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from a PDF using PyMuPDF (fitz).

    Returns an HTML fragment with one <h2>Page N</h2> per page followed by
    <p> blocks. Returns empty string on image-only / scanned PDFs.
    """
    import fitz  # PyMuPDF — lazy import keeps cold-start light

    parts: list[str] = ['<div data-source="upload-pdf">']
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for i, page in enumerate(doc, start=1):
                page_text = page.get_text("text") or ""
                page_paragraphs = _paragraphs_from_text(page_text)
                if not page_paragraphs:
                    continue
                parts.append(f"<h2>Page {i}</h2>")
                parts.extend(page_paragraphs)
    except Exception as exc:
        logger.error("extract_pdf_text failed: %s", exc)
        return ""
    parts.append("</div>")
    body = "\n".join(parts)
    # If the only content is the wrapper + opening, treat as empty
    if not re.search(r"<p>", body):
        return ""
    return body


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a DOCX using python-docx.

    Headings map to <h1>/<h2>/<h3>, default paragraphs to <p>, tables to
    plain <table> markup.
    """
    from docx import Document

    parts: list[str] = ['<div data-source="upload-docx">']
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        logger.error("extract_docx_text failed to open docx: %s", exc)
        return ""

    try:
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            style = (para.style.name if para.style else "") or ""
            style_lower = style.lower()
            if style_lower.startswith("heading 1") or style_lower == "title":
                parts.append(f"<h1>{escape(text)}</h1>")
            elif style_lower.startswith("heading 2"):
                parts.append(f"<h2>{escape(text)}</h2>")
            elif style_lower.startswith("heading 3"):
                parts.append(f"<h3>{escape(text)}</h3>")
            elif style_lower.startswith("heading"):
                parts.append(f"<h4>{escape(text)}</h4>")
            else:
                parts.append(f"<p>{escape(text)}</p>")

        for table in doc.tables:
            parts.append("<table>")
            for row in table.rows:
                parts.append("<tr>")
                for cell in row.cells:
                    cell_text = (cell.text or "").strip()
                    parts.append(f"<td>{escape(cell_text)}</td>")
                parts.append("</tr>")
            parts.append("</table>")
    except Exception as exc:
        logger.error("extract_docx_text failed during traversal: %s", exc)
        return ""

    parts.append("</div>")
    body = "\n".join(parts)
    if not re.search(r"<(p|h\d|table)>", body):
        return ""
    return body


def extract_html(file_bytes: bytes) -> str:
    """Sanitise uploaded HTML — strip scripts/styles/iframes and event handlers."""
    try:
        raw = file_bytes.decode("utf-8", errors="replace")
    except Exception as exc:
        logger.error("extract_html decode failed: %s", exc)
        return ""

    soup = BeautifulSoup(raw, "lxml")

    for tag in soup.find_all(list(_DANGEROUS_TAGS)):
        tag.decompose()

    for tag in soup.find_all(True):
        attrs_to_drop = [a for a in list(tag.attrs.keys()) if a.lower().startswith("on")]
        for a in attrs_to_drop:
            del tag.attrs[a]
        if tag.has_attr("href") and tag["href"].strip().lower().startswith("javascript:"):
            del tag["href"]
        if tag.has_attr("src") and tag["src"].strip().lower().startswith("javascript:"):
            del tag["src"]

    container = soup.body or soup
    inner = "".join(str(c) for c in container.contents) if hasattr(container, "contents") else str(container)
    inner = inner.strip()
    if not inner:
        return ""
    return f'<div data-source="upload-html">{inner}</div>'


def extract_any(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the right extractor based on filename extension.

    Returns an HTML fragment, or "" on failure / empty extraction. The caller
    treats an empty return as `extraction_failed` and falls back gracefully.
    """
    if not file_bytes:
        return ""
    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".pdf":
        return extract_pdf_text(file_bytes)
    if ext == ".docx":
        return extract_docx_text(file_bytes)
    if ext in (".html", ".htm"):
        return extract_html(file_bytes)
    raise ValueError(f"Unsupported extension '{ext}' for upload extraction")
