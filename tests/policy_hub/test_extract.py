"""Tests for policy_hub/extract.py — upload text extraction helpers."""

import io

import pytest

from policy_hub.extract import (
    extract_any,
    extract_docx_text,
    extract_html,
    extract_pdf_text,
)


class TestExtractHtml:
    def test_simple_html_preserved(self):
        html = b"<html><body><h1>Title</h1><p>Hello world</p></body></html>"
        out = extract_html(html)
        assert "Hello world" in out
        assert "<h1>Title</h1>" in out
        assert 'data-source="upload-html"' in out

    def test_strips_script_tags(self):
        html = b"<html><body><p>Safe</p><script>alert(1)</script></body></html>"
        out = extract_html(html)
        assert "<script" not in out
        assert "alert(1)" not in out
        assert "Safe" in out

    def test_strips_style_iframe_form(self):
        html = (
            b"<html><body>"
            b"<style>body{color:red}</style>"
            b"<iframe src='x'></iframe>"
            b"<form><input/></form>"
            b"<p>Kept</p>"
            b"</body></html>"
        )
        out = extract_html(html)
        assert "<style" not in out
        assert "<iframe" not in out
        assert "<form" not in out
        assert "Kept" in out

    def test_strips_event_handler_attrs(self):
        html = b'<html><body><div onclick="evil()">text</div></body></html>'
        out = extract_html(html)
        assert "onclick" not in out
        assert "evil()" not in out
        assert "text" in out

    def test_strips_javascript_urls(self):
        html = b'<html><body><a href="javascript:alert(1)">link</a></body></html>'
        out = extract_html(html)
        assert "javascript:" not in out

    def test_empty_html_returns_empty_string(self):
        out = extract_html(b"<html><body></body></html>")
        assert out == ""

    def test_html_without_body_tag_still_works(self):
        html = b"<p>Just a paragraph</p>"
        out = extract_html(html)
        assert "Just a paragraph" in out


class TestExtractDocx:
    """Build a minimal DOCX in-memory using python-docx, then round-trip through extract_docx_text."""

    @staticmethod
    def _make_docx(paragraphs):
        from docx import Document

        doc = Document()
        for style, text in paragraphs:
            if style:
                doc.add_paragraph(text, style=style)
            else:
                doc.add_paragraph(text)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_extracts_plain_paragraphs(self):
        docx_bytes = self._make_docx([
            (None, "First paragraph."),
            (None, "Second paragraph."),
        ])
        out = extract_docx_text(docx_bytes)
        assert "First paragraph." in out
        assert "Second paragraph." in out
        assert "<p>" in out
        assert 'data-source="upload-docx"' in out

    def test_maps_heading_styles(self):
        docx_bytes = self._make_docx([
            ("Heading 1", "Title"),
            ("Heading 2", "Subtitle"),
            (None, "Body content."),
        ])
        out = extract_docx_text(docx_bytes)
        assert "<h1>Title</h1>" in out
        assert "<h2>Subtitle</h2>" in out
        assert "<p>Body content.</p>" in out

    def test_empty_docx_returns_empty_string(self):
        from docx import Document
        buf = io.BytesIO()
        Document().save(buf)
        out = extract_docx_text(buf.getvalue())
        assert out == ""

    def test_corrupt_bytes_returns_empty_string(self):
        out = extract_docx_text(b"this is not a docx file at all")
        assert out == ""

    def test_html_special_chars_escaped(self):
        docx_bytes = self._make_docx([
            (None, "<script>evil</script>"),
        ])
        out = extract_docx_text(docx_bytes)
        assert "<script>" not in out
        assert "&lt;script&gt;" in out


class TestExtractPdf:
    """PDF extraction needs PyMuPDF — skip locally if not installed."""

    def setup_method(self):
        pytest.importorskip("fitz")

    def test_empty_bytes_returns_empty_string(self):
        out = extract_pdf_text(b"")
        assert out == ""

    def test_corrupt_pdf_returns_empty_string(self):
        out = extract_pdf_text(b"%PDF-1.4 garbage not a real pdf")
        assert out == ""


class TestExtractAny:
    def test_dispatches_html(self):
        out = extract_any(b"<p>hello</p>", "doc.html")
        assert "hello" in out

    def test_dispatches_htm_extension(self):
        out = extract_any(b"<p>hi</p>", "doc.HTM")
        assert "hi" in out

    def test_dispatches_docx(self):
        from docx import Document
        buf = io.BytesIO()
        d = Document()
        d.add_paragraph("Docx content.")
        d.save(buf)
        out = extract_any(buf.getvalue(), "test.docx")
        assert "Docx content." in out

    def test_unsupported_extension_raises(self):
        with pytest.raises(ValueError, match="Unsupported extension"):
            extract_any(b"some content", "doc.txt")

    def test_empty_bytes_returns_empty_string(self):
        out = extract_any(b"", "doc.html")
        assert out == ""
